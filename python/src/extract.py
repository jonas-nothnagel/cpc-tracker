"""
Document-to-target extraction.

Accepts a PDF, DOCX, or plain text file, extracts text, and uses the LLM to
identify high-level policy targets — the kind a human expert would pick out
for a cross-document coherence analysis.

Three-phase approach:
  0. (Optional) Relevance filter — discard chunks unlikely to contain targets.
  1. Per-chunk extraction of candidate targets (concurrent).
  2. Consolidation pass that deduplicates, merges, and filters candidates
     across the whole document into a final expert-quality list.

Usage:
    python -m src.extract --file path/to/document.pdf --doc-type NDC
"""

from __future__ import annotations

import argparse
import asyncio
import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .config import LLM_MODEL
from .extract_validation import (
    MatchCorpus,
    normalise_for_matching,
    validate_quotes_in_document,
)
from .footprint import append_event, electricity_zone
from .llm import (
    call_llm,
    call_llm_batch,
    call_llm_batch_detailed,
    get_footprint_tracker,
    set_language,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

MAX_CHUNK_CHARS = int(os.getenv("CPC_MAX_CHUNK_CHARS", "30000"))
OVERLAP_CHARS = int(os.getenv("CPC_OVERLAP_CHARS", "2000"))
MIN_TARGET_LENGTH = int(os.getenv("CPC_MIN_TARGET_LENGTH", "10"))

# Consolidation runs over page-ordered windows of at most this many
# candidates. One giant call re-emitting every candidate (PNSH: 79) blows
# past the completion-token budget, truncates mid-JSON, and used to zero
# the whole document.
CONSOLIDATE_WINDOW = int(os.getenv("CPC_CONSOLIDATE_WINDOW", "24"))

# Completion-token ceiling for the per-chunk extraction calls. Reasoning
# models spend completion tokens on reasoning before output, and a dense
# chunk (the NWRP Section-7 policy/strategy table) produces 19k+ chars of
# JSON: at the old 4000-token ceiling 4 of that document's 6 chunks came
# back truncated mid-JSON and extracted to zero. A ceiling, not a target —
# short responses cost the same as before.
EXTRACT_MAX_TOKENS = int(os.getenv("CPC_EXTRACT_MAX_TOKENS", "12000"))

# Documents whose full text fits this budget get the WHOLE document as
# activities-extraction context instead of quote-anchored windows. Why: the
# Sri Lanka fisheries policy lists its objectives up front, and they do not
# map 1-1 onto the action sections below; windows anchored on the objectives
# list missed most of the numbered actions the expert reviewer expected
# (2 Jul 2026 review). Windows remain for documents too large to pass whole.
ACTIVITY_FULLDOC_CHARS = int(os.getenv("CPC_ACTIVITY_FULLDOC_CHARS", "30000"))

# A PDF page with fewer extractable characters than this is treated as
# having no usable text layer (scanned/image page).
PDF_EMPTY_PAGE_CHARS = 50
# Above this fraction of empty pages the document warrants a partial-scan
# warning surfaced to the reviewer.
PDF_PARTIAL_EMPTY_RATIO = 0.3

# Meta-information about the most recent extract_from_text run (parse
# failures, consolidation fallbacks, document warnings). Reset per run and
# written to the `<output>.meta.json` sidecar by the CLI so the API route
# can surface warnings. Single-document-per-process usage only.
RUN_META: dict[str, Any] = {}

# ---------------------------------------------------------------------------
# Expected target counts by document type. WEAK, order-of-magnitude guidance
# only, derived from the early Mongolia corpus; real corpora range far wider
# (NDC: Mongolia 27, Sri Lanka 91, Côte d'Ivoire 181). The prompts present
# these as calibration that must never cap coverage.
# ---------------------------------------------------------------------------
EXPECTED_COUNTS: dict[str, str] = {
    "NDC": "15-30",
    "NBSAP": "15-25",
    "NAP": "10-20",
    "LDN": "5-15",
    "SECTORAL": "5-20",
}

# ---------------------------------------------------------------------------
# Few-shot examples, drawn verbatim from the human-curated corpora of several
# countries (Mongolia, Sri Lanka, Côte d'Ivoire) so the model calibrates on
# LEVEL and STRUCTURE, not on one country's phrasing.
#
# CONTAMINATION RULE: any document whose curated targets appear here must not
# serve as tier-1 gold in scripts/eval_manifest.json (currently excluded on
# these grounds: Mongolia FSS, Mongolia NBSAP, Sri Lanka NDC, Côte d'Ivoire NDC). Update the
# manifest exclusion note when changing the examples.
# ---------------------------------------------------------------------------
FEW_SHOT_EXAMPLES = """\
Here are examples of correctly extracted targets from real policy documents \
of several countries. Use these as a reference for the level of abstraction \
and specificity expected — document structures and phrasing vary widely by \
country; the LEVEL is what generalises.

NDC examples:
- "Protect natural forests, improve forest health, and enhance forest regeneration capacity"
- "Reduce Soil Organic Carbon loss by restoring 10% of heavily degraded rangelands within \
the forest-steppe (3%) and steppe (7%) ecological zones, moving land from Recovery Class IV \
(heavily degraded) to Recovery Class III (moderately degraded)."
- "Strengthen the health care system for early detection, warning, and response to climate \
change related threats to human health."
- "The agricultural sector will reduce 5.277 million tons of CO2 until 2030 and 9.78 million \
tons of CO2 until 2035 by adhering (national livestock herd to 50 million head) ecological \
carrying capacity of rangelands and improving livestock economic turnover through increasing \
export and meat supply to the domestic."

NBSAP examples:
- "By 2030, restore at least 30 percent of degraded ecosystems and improve ecological \
integrity and connectivity."
- "By 2030, 30% of the country's total area will be included in the special protected area \
network, representing major ecosystem types and strengthening ecological connectivity."
- "By 2030, reform subsidies and incentives that have negative impacts on biodiversity, \
while increasing environmentally positive (green) incentives."

NAP examples:
- "Minimize risks from climate change induced disasters and improve resilience"
- "Develop multi-hazard, impact-based early warning systems to enable quick and effective \
responses to extreme events."
- "Sustainably develop high-yield crops through climate-adapted technologies and soil \
conservation."

Sector-numbered targets (labels like "Forestry 1", "Transport 4") — short \
verbatim targets are valid, never pad them:
- "Sustainable management of forests and the restoration of other degraded lands, \
at least 32% by 2035"  (label: "Forestry 1")
- "Promote electric mobility"  (label: "Transport 4")
- "Promote a circular economy"  (label: "M24")

Tri-part structured target (some NDCs state each target as Goal / Target / \
Benchmark) — this is ONE target; keep the parts together, do not split them:
- "Goal: Preserve marine biodiversity and resources to enhance the resilience of \
coastal ecosystems against human and climate pressures.; Target: Achieve a \
coverage rate of at least 30% of sensitive marine areas by functional and \
sustainably managed Marine Protected Areas (MPAs) by 2035.; Benchmark: In 2020, \
less than 10% of sensitive marine areas in Côte d'Ivoire are covered by \
operational Marine Protected Areas (MPAs)."  (label: "A7.3")

Resolution-style document with a 3-level hierarchy (Section → Measure → Action):
Mongolia's Resolution 36 "Food Supply and Security Measures" has Section 2 broken \
into five numbered sub-sections (2/1 Legal environment, 2/2 Management, 2/3 Crop \
production, ...), each with multiple lettered items (а/, б/, в/, ...), and an annex \
table listing concrete sub-actions per item. The TARGET level is the MIDDLE level \
(the per-measure grouping, e.g. "1.1 Food/agri legislation drafts"), NOT every \
lettered item from the body and NOT every annex row. The annex sub-actions go into \
"activities" (when the schema supports it) or stay as multiple "sources" entries.
- Target: "To draft legislation aimed at establishing a legal framework for ensuring \
food supply and safety, to submit the draft law to the State Great Khural during \
the 2022 autumn and 2023 spring ordinary sessions; to reduce the country's reliance \
on imported food raw materials and products through tariff regulation, draft \
legislation aimed at creating a legal environment to increase the variety and \
volume of export-oriented food and agricultural products, and submit it to the \
State Great Khural by 2022"
  label: "1.1 Food/agri legislation drafts"
  (the annex sub-actions like "Draft the Law on Agriculture and submit it to the \
   State Great Khural", "Draft a bill to amend the Law on Food", "Draft a bill on \
   Agricultural Insurance" belong as activities under this one target, NOT as \
   separate targets — they implement 1.1, they are not independent objectives.)

Notice that these are policy-level OBJECTIVES, not:
- Individual implementation measures or activities (too granular)
- Background context or situation descriptions
- Indicator metrics in isolation (e.g. "Percentage of land under protection")
- Procedural text about stakeholder consultation or reporting
- Generic aspirational statements without a clear policy goal"""

# ---------------------------------------------------------------------------
# System and user prompts
# ---------------------------------------------------------------------------

EXTRACT_SYSTEM = """\
You are a senior policy analyst preparing structured data for a national \
policy coherence assessment. Your task is to extract the HIGH-LEVEL POLICY \
TARGETS from a section of a national policy document.

A "target" is a discrete policy-level objective, goal, or commitment that a \
government has set — the kind of statement that would appear in a summary \
table of a country's policy ambitions. Targets may be quantitative \
(e.g. "reduce emissions by 30% by 2030") or qualitative \
(e.g. "strengthen early warning systems for disaster risk").

{few_shot}

RULES:
1. Extract ONLY high-level policy targets/objectives/goals. NOT individual \
   measures, activities, indicators, background, or context.
2. When a document lists a "Goal" with sub-"Targets" or sub-"Measures", \
   extract at the TARGET level, not at the measure/activity level. Combine \
   closely related sub-points into one coherent target. For 3+ level \
   documents (e.g. Section → Measure → Action, or Resolution → numbered \
   sub-section → lettered item → annex action), the TARGET level is the \
   MIDDLE level. Lower-level items (specific actions, deliverables, draft \
   bills, individual procurement counts) belong under their parent measure \
   and NOT as separate targets.
3. VERBATIM CONTRACT — the "sourceText" field MUST be copied character-for-character \
   from the document, with no paraphrase, no reorder, no addition, no omission. \
   The "text" field MAY be a deterministic light cleanup of "sourceText", but only \
   the following changes are permitted:
     - drop a leading numeric/letter label ("Target 1.", "2.3", "a)") that is part \
       of the document's numbering, NOT semantic content;
     - join lines broken by PDF layout (replace internal "\\n" with a space);
     - collapse runs of whitespace into a single space;
     - drop trailing footnote markers like "[1]" or "¹";
   Anything beyond this — rewording, action-verb prefacing, summarising, expanding \
   abbreviations, adding context — is FORBIDDEN at this stage.
4. Do NOT extract background descriptions, context, definitions, procedural \
   text, or stakeholder lists. Statements assigning institutional roles and \
   responsibilities (which body administers, coordinates, or is custodian of \
   what) are governance descriptions, not policy targets — do not extract \
   them either.
5. If a section has no policy targets, return an empty array: []
6. As weak calibration only: a typical {doc_type} document contains \
   {expected_count} targets in total, but real counts vary widely by country \
   and document — some contain far more. Extract every DISTINCT policy target \
   this section actually contains; be selective about the LEVEL (Rules 1-2), \
   never about coverage.
7. If the document explicitly labels targets (e.g. "Target 1", "Goal 2.3"), \
   use those exact names as the label — do not rephrase or replace them.
8. Include all measurable details in the target text: percentages, quantities, \
   deadlines, years, baselines, or geographic scope. These details MUST already \
   be present in the source text — never add a number that is not in the document.
9. Do not extract each itemized line as its own target unless it is \
   independently tracked or measured. Group related items under a single target. \
   When you DO group, set "textCleanup" to "synthesis" and include each sub-item's \
   verbatim quote as a separate entry in "sources" (see schema below). When the \
   source document has an annex / action-plan table that enumerates concrete \
   actions per measure, those actions belong under their parent measure (as \
   `sources` entries or, downstream, as `activities`) — NOT as standalone targets.
10. If a table contains policy targets, extract from the full table context — \
    do not ignore targets just because they appear in tabular format.

Return a JSON array. Each object must have:
- "text": the policy target as it should be displayed and fed to downstream \
  classification (verbatim, cleaned, or synthesis per Rule 3 / Rule 9)
- "label": a short descriptive label (max 8 words; use document-provided \
  names verbatim when available)
- "sources": an array of one or more verbatim source spans. Each entry has:
    - "sourceText": the exact quote from the document, with no edits
    - "section": optional document-provided identifier (e.g. "Goal 2", "6.2.12")
- "textCleanup": one of "verbatim", "cleaned", or "synthesis":
    - "verbatim"  — text equals sources[0].sourceText with at most whitespace \
      normalisation
    - "cleaned"   — text is a deterministic light cleanup of a single source per \
      Rule 3 (label drop, line-join, whitespace collapse, footnote-marker drop)
    - "synthesis" — text combines multiple source entries; every concrete claim \
      in text must appear in at least one sources[].sourceText

Output valid JSON only — no markdown fences, no explanation."""

EXTRACT_USER = """\
Extract the high-level policy targets from the following section of a {doc_type} document:

---
{text}
---

Return a JSON array of extracted targets."""

# ---------------------------------------------------------------------------
# Relevance filter prompt
# ---------------------------------------------------------------------------

RELEVANCE_FILTER_SYSTEM = """\
You are filtering a section of a policy document for a target extraction \
pipeline. Your task is to classify whether this section is likely to contain \
policy targets, goals, or commitments.

Sections that are NOT relevant (return false):
- Table of contents, lists of abbreviations, acknowledgments, forewords
- Administrative/procedural text (stakeholder lists, meeting schedules)
- Institutional roles-and-responsibilities descriptions (which body \
administers or coordinates what) with no policy commitments or goals
- Background/context with no policy commitments or goals
- Annexes with only reference data, indicators, or methodology descriptions
- Repeated headers/footers, formatting artifacts, cover pages

Sections that ARE relevant (return true):
- Chapters describing policy goals, targets, or commitments
- Sections with measurable objectives, timelines, or deadlines
- Strategy/action sections with specific interventions
- Tables listing targets, measures, or actions
- Sections referencing specific policies, laws, or frameworks with goals

When uncertain, return true — it is better to keep a section than miss a target.

Return ONLY a JSON object: {"relevant": true} or {"relevant": false}
No explanation, no markdown fences."""

RELEVANCE_FILTER_USER = """\
Classify whether the following section of a {doc_type} document is likely to \
contain policy targets:

---
{text}
---"""

# ---------------------------------------------------------------------------
# Parallel-translation check prompt (multilingual documents)
# ---------------------------------------------------------------------------

PARALLEL_CHECK_SYSTEM = """\
You are checking the structure of a multilingual policy document for an \
extraction pipeline. Two blocks of pages from the SAME document are shown: \
BLOCK A is in the pipeline's working language; BLOCK B appears to be in a \
different language or script (possibly a legacy font encoding that renders \
as unreadable Latin characters).

Decide whether BLOCK B is a parallel translation of BLOCK A — the same \
policy content published again in another language — or whether it contains \
UNIQUE content that BLOCK A does not have.

Signals of a parallel translation: same document title and front matter, \
same section structure and numbering, the same tables, the same counts of \
numbered items, the same figures/percentages/years appearing in the same \
order.

When uncertain, answer {"parallel": false} — keeping a unique block is \
safe; skipping one loses content.

Return ONLY a JSON object: {"parallel": true} or {"parallel": false}
No explanation, no markdown fences."""

PARALLEL_CHECK_USER = """\
BLOCK A (working language, pages {a_pages}) — head and middle samples:
---
{a_sample}
---

BLOCK B (candidate parallel translation, pages {b_pages}) — head and middle samples:
---
{b_sample}
---

Is BLOCK B a parallel translation of BLOCK A?"""

# ---------------------------------------------------------------------------
# Consolidation prompt (second pass)
# ---------------------------------------------------------------------------

CONSOLIDATE_SYSTEM = """\
You are a senior policy analyst finalising the target list for a national \
policy coherence assessment. You have received candidate targets extracted \
from ONE PORTION of a {doc_type} document (consolidation runs per portion; \
other portions are handled separately).

Your task is to produce the CLEAN list of distinct policy targets for this \
portion by:
1. Removing duplicates or near-duplicates (keep the most complete version).
2. Merging overlapping targets that clearly refer to the same objective. When \
   you merge, you MUST preserve every "sources" entry from the merged \
   candidates — sources are the verbatim provenance trail and may not be \
   dropped, summarised, or rewritten.
3. Removing items that are NOT real policy targets (background text, \
   indicators in isolation, procedural statements, generic filler).
4. Keeping the list at the right level of abstraction — each item should be \
   a distinct policy objective suitable for cross-document coherence analysis.

Do NOT compress the list toward a target count: merge only genuine \
duplicates and overlapping statements of the same objective, and remove only \
items that are clearly not policy targets. Distinct objectives stay separate \
even when they are many. (For calibration only: a typical {doc_type} \
contains roughly {expected_count} targets across the WHOLE document; this \
portion may hold any share of that.)

Each candidate includes a "pageNumbers" array showing which pages it was \
found on. When merging candidates, combine their page numbers (union of all \
pages from merged items).

CLAIM-GROUNDING RULE for the merged "text" field:
- Every concrete claim in the merged text — numbers, percentages, deadlines, \
  named frameworks, geographic scope, specific verbs — MUST appear in at least \
  one of the merged "sources[].sourceText" entries.
- Do NOT add new claims, scope, or qualifiers that were not in any source.
- If two candidates contradict each other on a number, keep both candidates as \
  separate targets rather than inventing a synthesis.

When merging:
- "textCleanup" = "synthesis" if you combined wording from multiple sources, \
  "verbatim"/"cleaned" if you kept a single candidate's text.
- "sources" = the union of all merged candidates' source entries (no \
  deduplication of sourceText required — repeated quotes are evidence of \
  recurrence across sections).
- When candidates carry "textOriginal" / "labelOriginal" fields (non-English \
  documents), every output object MUST carry them too: merge the \
  original-language text under the same rules as "text", and never translate \
  "sources[].sourceText".

Output the final items in the same order as the input candidates (document \
order); merging keeps the position of the earliest merged candidate.

Return a JSON array. Each object must have:
- "text": the policy target (display + pipeline input)
- "label": a short descriptive label (max 8 words). When a candidate carries \
  a document-provided name ("Policy 1", "Goal 2.3", "Target 4"), keep that \
  name verbatim — never replace it with an invented phrase; when merging, \
  keep the kept candidate's label
- "pageNumbers": array of page numbers where this target appears
- "sources": array of {{ "sourceText": "...", "section": "..." }} entries — \
  preserved verbatim from input candidates, never invented
- "textCleanup": "verbatim" | "cleaned" | "synthesis"
- "textOriginal" and "labelOriginal": REQUIRED on every object whenever the \
  input candidates carry them (non-English documents) — the original-language \
  target text and label, merged under the same rules as "text"

Output valid JSON only — no markdown fences, no explanation."""

CONSOLIDATE_USER = """\
Here are {count} candidate targets extracted from a {doc_type} document. \
Produce the final consolidated list.

{candidates_json}"""

# ---------------------------------------------------------------------------
# Multi-language extraction prompt addendum
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Activities extraction prompt (Phase 3)
# ---------------------------------------------------------------------------

ACTIVITIES_SYSTEM = """\
You are a policy analyst extracting implementation details for policy targets.

Given a policy target and the surrounding source text where it was found, \
extract any explicitly listed sub-activities, measures, or actions that fall \
under this target.

RULES:
1. Extract ONLY sub-items that are explicitly listed in the source text \
   (numbered lists, bullet points, or clearly delineated measures).
2. VERBATIM CONTRACT — each activity's text MUST be copied directly from the \
   source. The only edits permitted are:
     - drop a leading numeric/letter label ("2.1", "a)") that is part of the \
       document's numbering, NOT semantic content;
     - join lines broken by PDF layout (replace internal "\\n" with a space);
     - collapse runs of whitespace into a single space.
   Do NOT paraphrase, summarise, embellish, or add scope. Do NOT invent or \
   infer activities that are not explicitly in the text.
3. If there are no explicit sub-activities, return an empty array: []
4. Preserve numbering from the source document when present (e.g. "2.1", "a)") \
   in the "section" field of each activity.

Return a JSON object whose "activities" array contains entries shaped like:
  {{
    "text":       "the activity as it should be displayed",
    "sourceText": "verbatim quote from the document — no edits",
    "section":    "optional source-document numbering (2.1, a, etc.)"
  }}

Output valid JSON only — no markdown fences, no explanation."""

ACTIVITIES_USER = """\
Target: {target_text}

Source text (from pages {pages}):
---
{context}
---

Extract any explicitly listed sub-activities, measures, or actions for this target."""

MULTILANG_ADDENDUM = """\

LANGUAGE NOTE: This document is written in {language_name} (not English).
- "sources[].sourceText" MUST remain the exact verbatim quote in the \
document's original language — never translate it.
- Return "text" and "label" as faithful ENGLISH translations of the target, \
preserving every number, percentage, date, and scope qualifier exactly.
- Additionally provide on each object:
  - "textOriginal": the target text in the original language, derived from \
the sources under the same verbatim/cleanup rules as "text"
  - "labelOriginal": the label in the original language
"""

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class PageSpan:
    """A span of text mapped to its source page number."""
    page: int
    text: str


@dataclass
class DocumentText:
    """Full document text with page-level provenance."""
    pages: list[PageSpan] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    def pages_for_range(self, start: int, end: int) -> list[int]:
        """Return page numbers that overlap with the character range [start, end)."""
        result: list[int] = []
        offset = 0
        for page_span in self.pages:
            page_end = offset + len(page_span.text) + 2  # +2 for \n\n separator
            if offset < end and page_end > start:
                if page_span.page not in result:
                    result.append(page_span.page)
            offset = page_end
        return result


@dataclass
class Chunk:
    """A text chunk with provenance metadata."""
    text: str
    pages: list[int]


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def _table_to_text(table) -> str:
    """Convert a pymupdf table to pipe-delimited text."""
    try:
        rows = table.extract()
        if not rows:
            return ""
        lines = []
        for row in rows:
            cells = [str(c or "").strip().replace("\n", " ") for c in row]
            lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception:
        return ""


def pdf_text_layer_stats(path: Path) -> dict[str, Any]:
    """Per-page extractable-text statistics for scanned-PDF detection.

    A page with fewer than PDF_EMPTY_PAGE_CHARS extractable characters is
    counted as empty (scanned/image-only). Callers decide policy: a fully
    empty document is an error, a partially empty one a reviewer warning.
    """
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf  # type: ignore[no-redef]

    with pymupdf.open(str(path)) as doc:
        page_chars = [len((page.get_text("text") or "").strip()) for page in doc]
    pages = len(page_chars)
    empty = sum(1 for c in page_chars if c < PDF_EMPTY_PAGE_CHARS)
    return {
        "pages": pages,
        "emptyPages": empty,
        "emptyRatio": round(empty / pages, 3) if pages else 0.0,
        "totalChars": sum(page_chars),
    }


def _extract_text_pdf(path: Path) -> list[PageSpan]:
    """Extract text from PDF with page-level tracking and table support."""
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf  # type: ignore[no-redef]

    doc = pymupdf.open(str(path))
    result: list[PageSpan] = []

    for i, page in enumerate(doc):
        page_num = i + 1
        text = page.get_text("text") or ""

        # Extract tables and append as structured text
        try:
            tables = page.find_tables()
            for table in tables:
                table_text = _table_to_text(table)
                if table_text:
                    text += f"\n\n[TABLE]\n{table_text}\n"
        except Exception:
            pass  # fall back to text-only if table extraction fails

        if text.strip():
            result.append(PageSpan(page=page_num, text=text.strip()))

    doc.close()
    return result


def _docx_table_to_text(table) -> str:
    """Convert a python-docx table to pipe-delimited text."""
    lines = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_text_docx(path: Path) -> list[PageSpan]:
    """Extract text from DOCX with table support and heading-based sections.

    DOCX has no reliable page boundaries, so every span carries page=0
    (unknown). Instead, the document is split at Heading-style paragraphs:
    each heading starts a new span with the heading kept as a "## " line, so
    the extraction prompts see section context and can populate
    sources[].section even without page numbers.
    """
    from docx import Document

    doc = Document(str(path))
    # O(1) element -> object lookups (the previous per-element linear scans
    # made large documents quadratic).
    para_by_el = {id(p._element): p for p in doc.paragraphs}
    table_by_el = {id(t._element): t for t in doc.tables}

    spans: list[PageSpan] = []
    parts: list[str] = []

    def flush() -> None:
        if parts:
            spans.append(PageSpan(page=0, text="\n\n".join(parts)))
            parts.clear()

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            para = para_by_el.get(id(element))
            text = (para.text if para is not None else element.text) or ""
            if not text.strip():
                continue
            is_heading = False
            if para is not None and para.style is not None:
                # style_id ("Heading1") is locale-invariant; the display name
                # is not ("Heading 1" / "Título 1" / "Titre 1"), so a
                # Spanish-authored DOCX would lose its sections on a
                # name-only check.
                style_id = getattr(para.style, "style_id", "") or ""
                style_name = para.style.name or ""
                is_heading = style_id.startswith("Heading") or style_name.startswith(
                    "Heading"
                )
            if is_heading:
                flush()
                parts.append(f"## {text.strip()}")
            else:
                parts.append(text.strip())
        elif tag == "tbl":
            table = table_by_el.get(id(element))
            if table is not None:
                table_text = _docx_table_to_text(table)
                if table_text:
                    parts.append(f"[TABLE]\n{table_text}")
    flush()

    if not spans:
        # Fallback: just paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return [PageSpan(page=0, text=full_text)] if full_text.strip() else []

    return spans


def extract_text(path: Path) -> DocumentText:
    """Extract text from a PDF, DOCX, or plain-text file with provenance."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = _extract_text_pdf(path)
    elif suffix == ".docx":
        pages = _extract_text_docx(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
        pages = [PageSpan(page=0, text=text)] if text.strip() else []

    return DocumentText(pages=pages)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def chunk_text(
    doc: DocumentText,
    max_chars: int = MAX_CHUNK_CHARS,
    overlap_chars: int = OVERLAP_CHARS,
) -> list[Chunk]:
    """Split document text into chunks with overlap, tracking source pages."""
    full_text = doc.full_text
    paragraphs = re.split(r"\n{2,}", full_text)
    chunks: list[Chunk] = []
    current = ""
    current_start = 0
    char_offset = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current) + len(para) + 2 > max_chars and current:
            # Save this chunk
            current_end = char_offset
            pages = doc.pages_for_range(current_start, current_end)
            chunks.append(Chunk(text=current, pages=pages))

            # Overlap: keep the tail of the current chunk
            if overlap_chars > 0 and len(current) > overlap_chars:
                overlap_text = current[-overlap_chars:]
                current_start = current_end - overlap_chars
                current = overlap_text + "\n\n" + para
            else:
                current_start = char_offset
                current = para
        else:
            current = f"{current}\n\n{para}" if current else para

        char_offset += len(para) + 2  # +2 for \n\n separator

    if current.strip():
        pages = doc.pages_for_range(current_start, char_offset)
        chunks.append(Chunk(text=current, pages=pages))

    return chunks


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------


def _language_samples(text: str, size: int = 5000) -> list[str]:
    """Head, middle, and tail samples for language detection.

    A single head sample misdetects multi-language documents: Sri Lankan
    policies open with Sinhala and Tamil sections before the English text,
    and legacy-font-encoded scripts read as Latin gibberish that lingua maps
    to arbitrary languages. Three positions give a majority vote a chance.
    """
    if len(text) <= size:
        return [text]
    mid_start = (len(text) - size) // 2
    return [text[:size], text[mid_start : mid_start + size], text[-size:]]


def _majority_language(codes: list[tuple[str, str]]) -> tuple[str, str, bool]:
    """Majority vote over per-sample detections.

    Returns (code, name, agreed). With no majority (all samples differ),
    falls back to the head sample's detection and reports agreed=False so
    the caller can warn: mixed-language documents deserve a human look and
    an explicit --language override.
    """
    counts: dict[str, int] = {}
    for code, _name in codes:
        counts[code] = counts.get(code, 0) + 1
    best_code = max(counts, key=lambda c: counts[c])
    if counts[best_code] > 1:
        name = next(name for code, name in codes if code == best_code)
        return best_code, name, counts[best_code] == len(codes)
    return codes[0][0], codes[0][1], False


def detect_language(text: str) -> tuple[str, str]:
    """Detect language of text. Returns (iso_code, language_name).

    Majority vote over head/middle/tail samples (multi-language documents);
    uses lingua-language-detector if available, falls back to langdetect,
    and ultimately defaults to English. Sample disagreement is logged and
    recorded in RUN_META as a document warning.
    """
    samples = _language_samples(text)

    def _finish(per_sample: list[tuple[str, str]]) -> tuple[str, str]:
        code, name, agreed = _majority_language(per_sample)
        if not agreed and len(per_sample) > 1:
            detail = [c for c, _ in per_sample]
            logger.warning(
                "Language detection disagrees across document sections %s — "
                "using %s. For multi-language documents pass --language "
                "explicitly.",
                detail, code,
            )
            RUN_META.setdefault("documentWarnings", []).append({
                "code": "MIXED_LANGUAGE_DETECTION",
                "samples": detail,
                "chosen": code,
            })
        return code, name

    try:
        from lingua import Language, LanguageDetectorBuilder
        detector = LanguageDetectorBuilder.from_all_languages().build()
        per_sample: list[tuple[str, str]] = []
        for sample in samples:
            detected = detector.detect_language_of(sample)
            if detected and detected != Language.ENGLISH:
                per_sample.append(
                    (detected.iso_code_639_1.name.lower(), detected.name.title())
                )
            else:
                per_sample.append(("en", "English"))
        return _finish(per_sample)
    except ImportError:
        pass

    try:
        from langdetect import detect
        per_sample = []
        for sample in samples:
            code = detect(sample)
            if code and code != "en":
                # langdetect returns ISO 639-1 codes
                per_sample.append((code, code.upper()))
            else:
                per_sample.append(("en", "English"))
        return _finish(per_sample)
    except ImportError:
        pass

    return "en", "English"


# ---------------------------------------------------------------------------
# Language blocks (parallel-translation handling)
#
# Some national policy PDFs publish the same content two or three times in
# sequential language blocks (observed: Sri Lanka NWRP 2023 — Sinhala, Tamil,
# then English, with the Sinhala/Tamil pages in legacy font encodings that
# extract as Latin gibberish). Extracting every block yields machine-
# translated duplicates of the working-language content and drowns out the
# document's own working-language text. The functions below segment a
# document into contiguous language blocks, deterministically; skipping a
# block additionally requires a numeric-fingerprint overlap AND an LLM
# confirmation, and is always surfaced as a document warning — never silent.
#
# Detection is deterministic on purpose: statistical detectors (lingua)
# return arbitrary languages at high confidence on legacy-font gibberish,
# so they must not drive per-page decisions.
# ---------------------------------------------------------------------------

# Unicode script ranges checked before any Latin-script heuristics. The
# Cyrillic → "mn" mapping reflects this pipeline's corpus (Mongolian policy
# documents), not a general claim about Cyrillic text.
_SCRIPT_RANGES: list[tuple[str, int, int]] = [
    ("si", 0x0D80, 0x0DFF),  # Sinhala
    ("ta", 0x0B80, 0x0BFF),  # Tamil
    ("mn", 0x0400, 0x04FF),  # Cyrillic (Mongolian in this corpus)
    ("ar", 0x0600, 0x06FF),  # Arabic
    ("hi", 0x0900, 0x097F),  # Devanagari
    ("zh", 0x4E00, 0x9FFF),  # CJK
]

# Stopword sets for the Latin-script languages the pipeline names. Density
# separates cleanly on real pages: English pages of the NWRP fixture score
# ~0.29-0.34, legacy-font gibberish 0.00-0.09.
_STOPWORDS: dict[str, frozenset[str]] = {
    "en": frozenset(
        "the of and to in for on with by is are be as that this from at or "
        "an it its will shall through".split()
    ),
    "es": frozenset(
        "el la los las de del y en para por con una un se que es al como su "
        "más este esta sobre entre".split()
    ),
    "fr": frozenset(
        "le la les des de du et en pour par avec une un dans que qui est au "
        "aux sur ce cette ainsi dont".split()
    ),
}
_STOPWORD_DENSITY_MIN = 0.12

_LANG_DISPLAY = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "mn": "Mongolian",
    "si": "Sinhala",
    "ta": "Tamil",
    "ar": "Arabic",
    "hi": "Hindi",
    "zh": "Chinese",
}

# Gates for skipping a non-working-language block as a parallel translation.
# All three must agree (size, numeric overlap, LLM confirmation); any
# uncertainty keeps the block.
_PARALLEL_MIN_PAGES = 5
_PARALLEL_MIN_CHARS = 10_000
_PARALLEL_MIN_NUMERIC_OVERLAP = 0.35


def _lang_display(code: str) -> str:
    return _LANG_DISPLAY.get(code, "an unidentified language")


def _classify_page_language(text: str) -> str:
    """Deterministic per-page language class.

    Returns an ISO-ish code from _SCRIPT_RANGES / _STOPWORDS, or "und" for
    pages with no reliable signal — short pages, numeric/TOC pages, and
    legacy-font gibberish all land in "und".
    """
    stripped = text.strip()
    if len(stripped) < 100:
        return "und"
    alpha = [ch for ch in stripped if ch.isalpha()]
    if not alpha:
        return "und"
    for code, lo, hi in _SCRIPT_RANGES:
        hits = sum(1 for ch in alpha if lo <= ord(ch) <= hi)
        if hits / len(alpha) > 0.5:
            return code
    tokens = re.findall(r"[a-zà-ÿ']+", stripped.lower())
    if len(tokens) < 20:
        return "und"
    best_code, best_density = "und", 0.0
    for code, words in _STOPWORDS.items():
        density = sum(1 for t in tokens if t in words) / len(tokens)
        if density > best_density:
            best_code, best_density = code, density
    return best_code if best_density >= _STOPWORD_DENSITY_MIN else "und"


@dataclass
class LanguageBlock:
    """A contiguous run of pages sharing one language class."""
    lang: str
    pages: list[PageSpan]

    @property
    def text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def chars(self) -> int:
        return sum(len(p.text) for p in self.pages)

    @property
    def page_range(self) -> list[int]:
        return [self.pages[0].page, self.pages[-1].page] if self.pages else []


def split_language_blocks(doc: DocumentText) -> list[LanguageBlock]:
    """Group contiguous pages into language blocks with run-length smoothing.

    A class change requires at least 2 consecutive pages of the new class;
    isolated single-page flips (photo pages, numeric tables, one-page
    summaries) absorb into the surrounding block. Documents that resolve to
    a single block behave exactly as before this function existed — the
    caller must gate all new behaviour on len(blocks) > 1.
    """
    pages = doc.pages
    if len(pages) < 2:
        return [LanguageBlock(lang=_classify_page_language(doc.full_text), pages=list(pages))]

    runs: list[list[Any]] = []  # [class | None, list[PageSpan]]
    for page in pages:
        cls = _classify_page_language(page.text)
        if runs and runs[-1][0] == cls:
            runs[-1][1].append(page)
        else:
            runs.append([cls, [page]])

    merged: list[list[Any]] = []
    for cls, run_pages in runs:
        if not merged:
            # A leading short run has no previous block; it adopts the class
            # of the first real run that follows (placeholder None).
            merged.append([cls if len(run_pages) >= 2 else None, list(run_pages)])
        elif len(run_pages) < 2 or cls == merged[-1][0]:
            merged[-1][1].extend(run_pages)
        elif merged[-1][0] is None:
            merged[-1][0] = cls
            merged[-1][1].extend(run_pages)
        else:
            merged.append([cls, list(run_pages)])

    # Adjacent same-class blocks can appear after absorption; fold them.
    folded: list[list[Any]] = []
    for cls, run_pages in merged:
        if folded and folded[-1][0] == cls:
            folded[-1][1].extend(run_pages)
        else:
            folded.append([cls, run_pages])

    if len(folded) == 1 and folded[0][0] is None:
        folded[0][0] = _classify_page_language(doc.full_text)

    # "und" runs are usually low-signal pages OF the surrounding block —
    # figure, table, and annex pages whose stopword density falls below the
    # threshold (observed: PNSH fragments into 8 runs this way). They stand
    # alone only when substantial relative to the largest known-language
    # block, which is the legacy-font parallel-translation signature
    # (NWRP: 111k chars of gibberish vs 50k of English); otherwise they
    # absorb into the neighbouring block, which restores the single-block
    # legacy path for ordinary documents.
    known_max = max(
        (
            sum(len(p.text) for p in run_pages)
            for cls, run_pages in folded
            if cls not in (None, "und")
        ),
        default=0,
    )
    absorbed: list[list[Any]] = []
    pending: list[PageSpan] = []  # leading und pages awaiting a block to join
    for cls, run_pages in folded:
        run_chars = sum(len(p.text) for p in run_pages)
        stands_alone = cls not in (None, "und") or not known_max or (
            run_chars >= 0.5 * known_max or len(run_pages) >= 15
        )
        if stands_alone:
            if pending:
                run_pages = pending + run_pages
                pending = []
            if absorbed and absorbed[-1][0] == cls:
                absorbed[-1][1].extend(run_pages)
            else:
                absorbed.append([cls, run_pages])
        elif absorbed:
            absorbed[-1][1].extend(run_pages)
        else:
            pending.extend(run_pages)
    if pending:
        if absorbed:
            absorbed[0][1][:0] = pending
        else:
            absorbed.append(["und", pending])

    return [LanguageBlock(lang=cls or "und", pages=run_pages) for cls, run_pages in absorbed]


_NUMERIC_TOKEN_RE = re.compile(r"\d+(?:[.,]\d+)*")


def _digit_jaccard(a: str, b: str) -> float:
    """Jaccard overlap of the numeric tokens in two texts.

    Parallel translations of the same policy content share their numbers
    (years, percentages, section counts) even when the scripts share nothing
    else. Measured on the NWRP fixture: Sinhala↔English 0.44, Tamil↔English
    0.76, vs 0.22 for an unrelated English document.
    """
    ta = set(_NUMERIC_TOKEN_RE.findall(a))
    tb = set(_NUMERIC_TOKEN_RE.findall(b))
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / len(ta | tb)


def _block_sample(text: str, head: int = 3000, mid: int = 1500) -> str:
    """Head + middle sample of a language block for the parallel check."""
    if len(text) <= head + mid:
        return text
    mid_start = (len(text) - mid) // 2
    return f"{text[:head]}\n[...]\n{text[mid_start : mid_start + mid]}"


def _parse_parallel(raw: str) -> bool:
    """Parse the parallel-check response. Defaults to False (keep the block)."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data.get("parallel") is True
    except json.JSONDecodeError:
        pass
    return False


async def _select_language_blocks(
    blocks: list[LanguageBlock],
    working_lang: str,
) -> list[LanguageBlock]:
    """Decide which language blocks of a multi-block document to extract.

    Working-language blocks are always kept. A non-working block is skipped
    ONLY as a confirmed parallel translation: it must be substantial
    (>= _PARALLEL_MIN_PAGES pages or _PARALLEL_MIN_CHARS chars), share the
    working block's numeric fingerprint (>= _PARALLEL_MIN_NUMERIC_OVERLAP),
    AND be confirmed by the LLM parallel check. Anything less is kept and
    extracted as today, with a warning. Every decision lands in
    RUN_META["languageBlocks"]; skips additionally emit a
    PARALLEL_TRANSLATION_SKIPPED document warning.
    """
    working = [b for b in blocks if b.lang == working_lang]
    block_meta: list[dict[str, Any]] = []
    if not working:
        RUN_META.setdefault("documentWarnings", []).append({
            "code": "MIXED_LANGUAGE_CONTENT",
            "detail": (
                f"document splits into {len(blocks)} language blocks but none "
                f"matches the working language '{working_lang}'; all blocks kept"
            ),
        })
        RUN_META["languageBlocks"] = [
            {"lang": b.lang, "pages": b.page_range, "chars": b.chars, "action": "kept"}
            for b in blocks
        ]
        return blocks

    working_text = "\n\n".join(b.text for b in working)

    candidates: list[tuple[int, float]] = []  # (block index, numeric overlap)
    calls: list[dict[str, Any]] = []
    for i, block in enumerate(blocks):
        if block.lang == working_lang:
            continue
        substantial = (
            len(block.pages) >= _PARALLEL_MIN_PAGES
            or block.chars >= _PARALLEL_MIN_CHARS
        )
        overlap = _digit_jaccard(block.text, working_text) if substantial else 0.0
        if substantial and overlap >= _PARALLEL_MIN_NUMERIC_OVERLAP:
            candidates.append((i, overlap))
            calls.append({
                "system": PARALLEL_CHECK_SYSTEM,
                "user": PARALLEL_CHECK_USER.format(
                    a_pages="-".join(str(p) for p in working[0].page_range),
                    a_sample=_block_sample(working_text),
                    b_pages="-".join(str(p) for p in block.page_range),
                    b_sample=_block_sample(block.text),
                ),
                "max_tokens": 50,
            })

    confirmed: dict[int, float] = {}
    if calls:
        results = await call_llm_batch(
            calls, cache_namespace="parallel_check", desc="Parallel-translation check"
        )
        for (i, overlap), raw in zip(candidates, results):
            if _parse_parallel(raw):
                confirmed[i] = overlap

    kept: list[LanguageBlock] = []
    for i, block in enumerate(blocks):
        meta: dict[str, Any] = {
            "lang": block.lang,
            "pages": block.page_range,
            "chars": block.chars,
        }
        if block.lang == working_lang:
            meta["action"] = "working"
            kept.append(block)
        elif i in confirmed:
            meta["action"] = "skipped-parallel"
            meta["numericOverlap"] = round(confirmed[i], 2)
            RUN_META.setdefault("documentWarnings", []).append({
                "code": "PARALLEL_TRANSLATION_SKIPPED",
                "pages": block.page_range,
                "langGuess": block.lang,
                "numericOverlap": round(confirmed[i], 2),
                "workingLanguage": working_lang,
            })
            logger.info(
                "Skipping pages %s (%s): confirmed parallel translation of the "
                "%s block (numeric overlap %.2f)",
                block.page_range, block.lang, working_lang, confirmed[i],
            )
        else:
            meta["action"] = "kept"
            kept.append(block)
            RUN_META.setdefault("documentWarnings", []).append({
                "code": (
                    "UNKNOWN_LANGUAGE_BLOCK"
                    if block.lang == "und"
                    else "MIXED_LANGUAGE_CONTENT"
                ),
                "pages": block.page_range,
                "langGuess": block.lang,
            })
        block_meta.append(meta)

    RUN_META["languageBlocks"] = block_meta
    return kept


# ---------------------------------------------------------------------------
# JSON parsing
# ---------------------------------------------------------------------------


_VALID_CLEANUPS = {"verbatim", "cleaned", "synthesis"}

# Section anchors the LLM claims must look like something a document would
# actually provide: numbering ("7", "6.2.12", "II/3"), a section word with
# optional numbering ("Goal 2", "Annex III", "Chapter 4: Adaptation"), or a
# plausible natural-language heading (DOCX extraction deliberately feeds
# heading names as sections). Everything else — echoed [TABLE] markers,
# legacy-font gibberish like "nfhs;if" — is dropped, never invented.
_SECTION_NUMBERING_RE = re.compile(
    r"^[\divxlc]+(?:[./-][\divxlca-z]+)*\.?$", re.IGNORECASE
)
_SECTION_WORD_RE = re.compile(
    r"^(?:goal|policy|policies|strategy|strategies|target|objective|section|"
    r"chapter|annex|appendix|part|article|measure|action|priority|pillar|"
    r"axis|outcome|table)\b[\w\s.:()/-]{0,60}$",
    re.IGNORECASE,
)
_VOWEL_RE = re.compile(r"[aeiouáéíóúàèìòùâêîôûäëïöü]", re.IGNORECASE)


def _valid_section_part(part: str) -> bool:
    if _SECTION_NUMBERING_RE.match(part) or _SECTION_WORD_RE.match(part):
        return True
    # Plausible natural-language heading: mostly letters, contains a vowel,
    # no stray symbols, not a two-letter fragment.
    if len(part) < 4 or len(part) > 80 or "%" in part:
        return False
    compact = part.replace(" ", "")
    letters = sum(1 for c in compact if c.isalpha())
    return bool(_VOWEL_RE.search(part)) and letters >= 0.6 * max(len(compact), 1)


def _clean_section(raw: Any) -> str | None:
    """Validate/normalise an LLM-provided section anchor. None = drop it."""
    if raw is None:
        return None
    s = re.sub(r"\s+", " ", str(raw)).strip()
    if not s:
        return None
    kept = []
    for part in re.split(r"[;,]", s):
        part = part.strip().strip("[]")
        # The literal TABLE token is the text extractors' own [TABLE] marker
        # echoed back — exact-token only, so "Table 3" survives.
        if not part or part.upper() == "TABLE":
            continue
        if _valid_section_part(part):
            kept.append(part)
    if not kept:
        if s:
            dropped = RUN_META.setdefault("sectionsDropped", {"count": 0, "examples": []})
            dropped["count"] += 1
            if len(dropped["examples"]) < 10 and s not in dropped["examples"]:
                dropped["examples"].append(s[:40])
        return None
    return "; ".join(kept)[:120]


def _parse_sources(item: dict[str, Any]) -> list[dict[str, Any]]:
    """Extract a normalised sources[] list from an LLM extraction item.

    Accepts either the new schema ("sources": [{sourceText, section?}, ...]) or
    the legacy single-source schema ("sourceText": "...") and returns a list of
    {sourceText, section?} dicts. Empty-string sources are dropped; section
    anchors are validated by _clean_section.
    """
    raw = item.get("sources")
    out: list[dict[str, Any]] = []
    if isinstance(raw, list):
        for src in raw:
            if isinstance(src, dict):
                txt = str(src.get("sourceText", "")).strip()
                if not txt:
                    continue
                entry: dict[str, Any] = {"sourceText": txt}
                section = _clean_section(src.get("section"))
                if section:
                    entry["section"] = section
                out.append(entry)
            elif isinstance(src, str) and src.strip():
                out.append({"sourceText": src.strip()})
    elif isinstance(item.get("sourceText"), str) and item["sourceText"].strip():
        out.append({"sourceText": item["sourceText"].strip()})
    return out


def _parse_json_array(
    raw: str,
    min_length: int = MIN_TARGET_LENGTH,
) -> list[dict[str, Any]]:
    """Parse LLM JSON response, handling common formatting issues."""
    return _parse_json_array_status(raw, min_length=min_length)[0]


def _salvage_truncated_array(raw: str) -> list[Any] | None:
    """Recover the complete leading objects of a truncated JSON array.

    A response cut off at the completion-token ceiling ends mid-object,
    usually mid-string (observed on the NWRP Section-7 table chunk: 19k
    chars ending inside a sourceText). Every object that closed cleanly is
    still valid — keep those instead of dropping the whole chunk. Returns
    None when nothing is recoverable.
    """
    start = raw.find("[")
    if start < 0:
        return None
    depth = 0
    in_str = False
    esc = False
    last_complete = -1
    for i in range(start, len(raw)):
        ch = raw[i]
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
        elif ch in "[{":
            depth += 1
        elif ch in "]}":
            depth -= 1
            if depth == 1 and ch == "}":
                last_complete = i
    if last_complete < 0:
        return None
    try:
        items = json.loads(raw[start : last_complete + 1] + "]")
    except json.JSONDecodeError:
        return None
    return items if isinstance(items, list) and items else None


def _parse_json_array_status(
    raw: str,
    min_length: int = MIN_TARGET_LENGTH,
    salvage_truncated: bool = False,
) -> tuple[list[dict[str, Any]], bool, bool]:
    """Like _parse_json_array but also reports HOW parsing went.

    Returns (items, parse_ok, truncated). A legitimately empty array ("[]")
    returns ([], True, False); a malformed response returns ([], False,
    False). Callers that would otherwise silently discard upstream work
    (consolidation) must check parse_ok and fall back instead of treating
    garbage as "no targets".

    With salvage_truncated=True, a response that fails to parse is treated
    as a truncation and its complete leading objects are recovered
    (items, True, True). Only safe where partial recovery beats total loss
    — per-chunk extraction. NEVER enable it for consolidation: candidates
    beyond the cut would be silently lost, whereas the parse-failure
    fallback keeps the whole un-consolidated window.
    """
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

    truncated = False
    try:
        items = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        items = None
        if match:
            try:
                items = json.loads(match.group())
            except json.JSONDecodeError:
                items = None
        if items is None and salvage_truncated:
            items = _salvage_truncated_array(raw)
            if items is not None:
                truncated = True
                logger.warning(
                    "Extraction response truncated mid-JSON — salvaged %d "
                    "complete item(s); the chunk's tail items are lost",
                    len(items),
                )
        if items is None:
            if raw:
                logger.warning("Could not parse extraction response as JSON")
            # Empty responses (e.g. content-filtered calls) are NOT a
            # trustworthy "no targets" signal either.
            return [], False, False

    if not isinstance(items, list):
        return [], False, False

    valid = []
    for item in items:
        if isinstance(item, dict) and "text" in item:
            text = str(item["text"]).strip()
            if len(text) < min_length:
                logger.debug(f"Skipping short target ({len(text)} chars): {text[:50]}")
                continue
            if min_length < 20 <= len(text):
                pass  # normal
            elif len(text) < 20:
                logger.info(f"Keeping short target ({len(text)} chars): {text[:50]}")

            entry: dict[str, Any] = {
                "text": text,
                "label": str(item.get("label", "")).strip()[:80]
                    or f"Target {len(valid) + 1}",
            }

            # Preserve page numbers if present
            if "pageNumbers" in item and isinstance(item["pageNumbers"], list):
                entry["pageNumbers"] = item["pageNumbers"]

            # Original-language fields (English-first schema): text/label are
            # English; textOriginal/labelOriginal carry the source language.
            for tfield in ("textOriginal", "labelOriginal"):
                if tfield in item and item[tfield]:
                    entry[tfield] = str(item[tfield]).strip()
            # Legacy schema (pre English-first cached responses): original in
            # text, English in text_eng. Invert into the canonical shape.
            if item.get("text_eng") and "textOriginal" not in entry:
                entry["textOriginal"] = entry["text"]
                entry["labelOriginal"] = entry["label"]
                entry["text"] = str(item["text_eng"]).strip()
                if item.get("label_eng"):
                    entry["label"] = str(item["label_eng"]).strip()[:80]

            # Provenance: sources[] (verbatim quotes) + textCleanup
            sources = _parse_sources(item)
            if sources:
                entry["sources"] = sources
            cleanup = str(item.get("textCleanup", "")).strip().lower()
            if cleanup in _VALID_CLEANUPS:
                entry["textCleanup"] = cleanup
            elif sources:
                # Default: assume verbatim if a single source matches the
                # original-language text (fall back to `text` for English
                # documents) after whitespace normalisation; otherwise default
                # to "cleaned" so downstream consumers can still surface a
                # provenance badge.
                compare = entry.get("textOriginal") or entry["text"]
                norm_text = re.sub(r"\s+", " ", compare).strip()
                norm_src = re.sub(r"\s+", " ", sources[0]["sourceText"]).strip()
                entry["textCleanup"] = "verbatim" if norm_text == norm_src else "cleaned"

            valid.append(entry)
    return valid, True, truncated


def _expected_count(source_document: str, doc_type: str) -> str:
    """Typical-target-count guidance for the prompts.

    source_document is the canonical type code (the wizard sends the same
    value for both args); doc_type is only a prompt label and may be a
    free-form override on the CLI, so it is the fallback, not the key.
    """
    return (
        EXPECTED_COUNTS.get(source_document.upper())
        or EXPECTED_COUNTS.get(doc_type.upper())
        or "10-30"
    )


# ---------------------------------------------------------------------------
# Deterministic candidate dedup (chunk-overlap artifacts)
# ---------------------------------------------------------------------------


def _merge_candidate(kept: dict[str, Any], dup: dict[str, Any]) -> None:
    """Fold a near-duplicate candidate into the kept one.

    If the duplicate's text is longer (more complete), its payload wins;
    sources and pageNumbers are unioned either way.
    """
    if len(str(dup.get("text", ""))) > len(str(kept.get("text", ""))):
        for field_ in ("text", "label", "textCleanup", "text_eng", "label_eng",
                       "textOriginal", "labelOriginal"):
            if dup.get(field_):
                kept[field_] = dup[field_]
    seen_sources = {
        (s.get("sourceText"), s.get("section")) for s in kept.get("sources") or []
    }
    for src in dup.get("sources") or []:
        key = (src.get("sourceText"), src.get("section"))
        if key not in seen_sources:
            seen_sources.add(key)
            kept.setdefault("sources", []).append(src)
    pages = set(kept.get("pageNumbers") or []) | set(dup.get("pageNumbers") or [])
    if pages:
        kept["pageNumbers"] = sorted(pages)


def dedupe_candidates(candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Merge exact and near-duplicate candidates deterministically.

    The 2000-char chunk overlap re-extracts targets that straddle a chunk
    boundary, so the same target routinely appears twice with near-identical
    wording. Near-dup = normalised texts are equal, one contains the other
    at comparable length, or their sequence ratio is >= 0.92. Runs before
    (and independently of) the LLM consolidation pass, so small documents
    that skip consolidation still get deduplicated.
    """
    from difflib import SequenceMatcher

    merged: list[dict[str, Any]] = []
    norms: list[str] = []
    for cand in candidates:
        norm = normalise_for_matching(str(cand.get("text", "")))
        hit_idx = None
        for i, kept_norm in enumerate(norms):
            if norm == kept_norm:
                hit_idx = i
                break
            # Full containment of a long, specific policy string means the
            # same commitment (one chunk simply saw more context); sources
            # and pages are unioned, so nothing is lost by merging.
            shorter, longer = sorted((norm, kept_norm), key=len)
            if len(shorter) >= 40 and shorter in longer:
                hit_idx = i
                break
            sm = SequenceMatcher(None, norm, kept_norm)
            if (
                sm.real_quick_ratio() >= 0.92
                and sm.quick_ratio() >= 0.92
                and sm.ratio() >= 0.92
            ):
                hit_idx = i
                break
        if hit_idx is None:
            merged.append(cand)
            norms.append(norm)
        else:
            _merge_candidate(merged[hit_idx], cand)
            norms[hit_idx] = normalise_for_matching(str(merged[hit_idx].get("text", "")))
    if len(merged) < len(candidates):
        logger.info(
            "Deterministic dedup: %d -> %d candidates", len(candidates), len(merged)
        )
    return merged


# ---------------------------------------------------------------------------
# Claim-grounding validator
# ---------------------------------------------------------------------------

# Regexes for "concrete claims" that, if present in display `text`, must also
# appear in at least one source span. Kept conservative — false positives are
# acceptable (the entry gets flagged for review), but we never want to miss a
# fabricated number / year / unit.
#
# Each pattern returns the raw match; we normalise (strip thousands separators,
# lowercase units) before set-membership.
_CLAIM_PATTERNS: list[tuple[str, "re.Pattern[str]"]] = [
    ("percent", re.compile(r"\b\d{1,3}(?:[.,]\d+)?\s?%")),
    ("year", re.compile(r"\b(?:19|20)\d{2}\b")),
    # Quantities with units we care about in policy targets. English `text`
    # is compared against original-language sources (es/mn/fr corpora), so
    # the alternation carries the unit spellings observed in those corpora;
    # _normalise_claim strips separators, so only the unit word must match a
    # variant here.
    (
        "quantity",
        re.compile(
            r"\b\d[\d.,]*\s?(?:tCO2e|tCO2eq|tCO₂e|tCO₂eq|MtCO2|"
            r"million tons?|thousand tons?|tons?|tonnes?|ha|km2|km²|"
            r"hectares?|GW|MW|kWh|"
            r"hect[áa]reas?|millones\s+de\s+toneladas?|miles\s+de\s+toneladas?|"
            r"toneladas?|"
            r"millions?\s+de\s+tonnes?|milliers?\s+de\s+tonnes?|"
            r"га|тонн|сая\s+тонн|мянган?\s+тонн)\b",
            re.IGNORECASE,
        ),
    ),
    # Plain large numbers (>= 4 digits when no separator, or any with separator).
    # Skips bare years; the year pattern above already covers those.
    ("number", re.compile(r"\b\d{1,3}(?:[.,]\d{3})+(?:[.,]\d+)?\b")),
]


def _normalise_claim(token: str) -> str:
    """Normalise a claim token for set comparison.

    Strips whitespace, lowercases, and collapses both `,` and `.` thousands
    separators so that "5,277", "5.277" and "5277" compare as equal — the
    LLM occasionally re-renders the same value with a different separator
    than the source, and we don't want that to false-flag a grounded claim.
    Both text and source go through the same normalisation, so decimals
    like "1.5%" still match each other (they normalise to "15%" on both
    sides — equivalent for set membership).
    """
    s = re.sub(r"\s+", "", token).lower()
    return s.replace(",", "").replace(".", "")


def _extract_claims(text: str) -> set[str]:
    """Return the set of normalised concrete claims found in text."""
    claims: set[str] = set()
    for _kind, pattern in _CLAIM_PATTERNS:
        for match in pattern.findall(text):
            claims.add(_normalise_claim(match))
    return claims


_NUM_PREFIX_RE = re.compile(r"^\d+")


def _missing_claims(claims: set[str], src_claims: set[str]) -> list[str]:
    """Claims not grounded in the source claims.

    Exact set-membership first. For non-percent claims, a numeric-magnitude
    fallback grounds e.g. "1500hectares" (English display text) against
    "1500hectáreas" (Spanish source): unit words are language-specific, the
    magnitude is not. Tradeoff, accepted: a same-magnitude unit swap
    (tons vs hectares) passes the fallback; the quote-in-document validator
    and the human review remain the backstop.
    """
    src_numeric = {
        m.group() for c in src_claims if (m := _NUM_PREFIX_RE.match(c))
    }
    missing = []
    for c in sorted(claims):
        if c in src_claims:
            continue
        m = _NUM_PREFIX_RE.match(c)
        if m and not c.endswith("%") and m.group() in src_numeric:
            continue
        missing.append(c)
    return missing


def validate_claim_grounding(target: dict[str, Any]) -> list[str]:
    """Return a list of unsourced claims found in the target's display text.

    Empty list means the target is well-grounded. Checks the English `text`
    and, when present, the original-language `textOriginal` against the
    verbatim sources. Run on every extracted target after consolidation;
    surface non-empty results as warnings so the pipeline does not silently
    emit fabricated numbers.
    """
    text = str(target.get("text", ""))
    if not text:
        return []
    sources = target.get("sources") or []
    src_blob = " ".join(
        str(s.get("sourceText", "")) for s in sources if isinstance(s, dict)
    )
    if not src_blob:
        # No sources at all — caller decides how strict to be; the validator
        # only flags claims, not missing-sources policy.
        return []

    src_claims = _extract_claims(src_blob)
    missing = set(_missing_claims(_extract_claims(text), src_claims))
    if target.get("textOriginal"):
        missing |= set(
            _missing_claims(
                _extract_claims(str(target["textOriginal"])), src_claims
            )
        )
    return sorted(missing)


def _relevance_sample(text: str, limit: int = 8000) -> str:
    """Representative sample of a chunk for the relevance filter.

    Judging only the head of a ~30k chunk misclassifies chunks whose
    substantive content starts after boilerplate: the Sri Lanka NAP goals
    sat on page 4 behind three pages of cover/foreword, and the whole
    pages-1-8 chunk was dropped as irrelevant. Head + middle + tail slices
    give the filter sight of the entire chunk at the same token budget.
    """
    if len(text) <= limit:
        return text
    head = text[: limit // 2]
    quarter = limit // 4
    mid_start = (len(text) - quarter) // 2
    mid = text[mid_start : mid_start + quarter]
    tail = text[-quarter:]
    return f"{head}\n[...]\n{mid}\n[...]\n{tail}"


def _parse_relevance(raw: str) -> bool:
    """Parse relevance filter response. Defaults to True (conservative)."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data.get("relevant", True)
    except json.JSONDecodeError:
        pass
    # Default to relevant when uncertain
    return True


# ---------------------------------------------------------------------------
# Phase 3: Activities extraction
# ---------------------------------------------------------------------------


def _activity_context(
    target: dict[str, Any],
    doc: DocumentText,
    norm_chunks: list[tuple["Chunk", str]],
    page_to_chunks: dict[int, list[str]],
) -> str:
    """Assemble the source context for one target's activities call.

    Whole document when it fits ACTIVITY_FULLDOC_CHARS; otherwise
    quote-anchored windows; otherwise (quotes unlocatable) the chunks
    covering the target's pages, truncated to the window budget.
    """
    full_text = doc.full_text
    if len(full_text) <= ACTIVITY_FULLDOC_CHARS:
        return full_text

    context_parts = _quote_context_windows(target, norm_chunks)
    if not context_parts:
        seen: set[int] = set()
        for page in target.get("pageNumbers", []):
            for chunk_text in page_to_chunks.get(page, []):
                h = hash(chunk_text)
                if h not in seen:
                    seen.add(h)
                    context_parts.append(chunk_text)

    context = "\n\n---\n\n".join(context_parts)
    if len(context) > 12000:
        context = context[:12000].rsplit("\n\n", 1)[0]  # truncate at paragraph boundary
    return context


def _quote_context_windows(
    target: dict[str, Any],
    norm_chunks: list[tuple["Chunk", str]],
    *,
    before: int = 1500,
    after: int = 4500,
    max_windows: int = 3,
) -> list[str]:
    """Context windows around the target's source-quote occurrences.

    Policy documents elaborate a statement into its action list immediately
    AFTER the statement text, so windows extend mostly forward. A target
    whose sources were merged from a summary list AND a detailed section
    gets one window per occurrence, so the detailed section is always in
    context. Offsets are approximated by position ratio between normalised
    and original chunk text (normalisation removes characters roughly
    uniformly); a few hundred chars of drift is irrelevant at window size.
    """
    windows: list[str] = []
    seen_spans: list[tuple[int, int]] = []
    for src in target.get("sources") or []:
        if len(windows) >= max_windows:
            break
        if not isinstance(src, dict):
            continue
        quote = str(src.get("sourceText", "")).strip()
        if not quote:
            continue
        norm_quote = normalise_for_matching(quote)
        if not norm_quote:
            continue
        for ci, (chunk, norm) in enumerate(norm_chunks):
            pos = norm.find(norm_quote)
            if pos < 0:
                continue
            center = int(len(chunk.text) * (pos / max(len(norm), 1)))
            start = max(0, center - before)
            end = min(len(chunk.text), center + after)
            if any(c == ci and abs(start - s) < 1000 for c, s in seen_spans):
                break  # near-duplicate of a window we already took
            seen_spans.append((ci, start))
            windows.append(chunk.text[start:end])
            break
    return windows


async def _extract_activities(
    targets: list[dict[str, Any]],
    doc: DocumentText,
    chunks: list[Chunk],
    is_english: bool,
    lang_name: str,
) -> list[dict[str, Any]]:
    """For each target, extract sub-activities from the surrounding source text."""
    # Build a lookup from page number to chunk text for context retrieval
    page_to_chunks: dict[int, list[str]] = {}
    for chunk in chunks:
        for page in chunk.pages:
            page_to_chunks.setdefault(page, []).append(chunk.text)

    # Context selection, by document size:
    # - documents that fit ACTIVITY_FULLDOC_CHARS whole: pass the ENTIRE
    #   document. Objectives lists up front often do not map 1-1 onto the
    #   action sections below, so any anchoring heuristic misses actions.
    # - larger documents: windows anchored on each target's verbatim quote
    #   occurrences. Page-chunk dumps failed two ways on real documents: a
    #   consolidated target whose first quote sits in an early summary chunk
    #   got 12k of front matter (truncation never reached the section that
    #   lists its actions), and whole chunks exceed the 12k budget anyway.
    norm_chunks = [(chunk, normalise_for_matching(chunk.text)) for chunk in chunks]

    activity_calls = []
    for target in targets:
        pages = target.get("pageNumbers", [])
        context = _activity_context(target, doc, norm_chunks, page_to_chunks)
        pages_str = ", ".join(str(p) for p in pages) if pages else "unknown"

        system = ACTIVITIES_SYSTEM
        if not is_english:
            system += (
                f"\n\nThe source text is in {lang_name}. Return each "
                'activity\'s "text" as a faithful ENGLISH translation '
                "(preserve every number, percentage, and date exactly); "
                '"sourceText" MUST remain the verbatim original-language '
                "quote — never translate it."
            )

        activity_calls.append({
            "system": system,
            "user": ACTIVITIES_USER.format(
                target_text=target["text"],
                pages=pages_str,
                context=context,
            ),
            # Reasoning models spend completion tokens on reasoning before
            # output; 2000 starved long activity lists into empty responses.
            "max_tokens": 6000,
        })

    raw_results = await call_llm_batch(
        activity_calls,
        cache_namespace="extract_activities",
        desc="Activities extraction",
    )

    for target, raw in zip(targets, raw_results):
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        try:
            data = json.loads(raw)
            activities = data.get("activities", [])
        except (json.JSONDecodeError, AttributeError):
            continue  # No activities — that's fine

        if not isinstance(activities, list) or not activities:
            continue

        # Normalise: accept either ["str", ...] (legacy) or [{text, sourceText, section?}].
        # Stored as a newline-joined display string in `activities` (existing
        # consumer contract) plus a structured `activitySources` array carrying
        # the verbatim quotes, so downstream UI can show provenance per activity.
        display_lines: list[str] = []
        structured: list[dict[str, Any]] = []
        for a in activities:
            if isinstance(a, dict):
                txt = str(a.get("text", "")).strip()
                src_txt = str(a.get("sourceText", "")).strip()
                section = _clean_section(a.get("section"))
                if not txt and src_txt:
                    txt = src_txt
                if not txt:
                    continue
                display_lines.append(txt)
                entry: dict[str, Any] = {"text": txt}
                if src_txt:
                    entry["sourceText"] = src_txt
                if section:
                    entry["section"] = section[:60]
                structured.append(entry)
            elif isinstance(a, str) and a.strip():
                display_lines.append(a.strip())
                structured.append({"text": a.strip()})

        if display_lines:
            target["activities"] = "\n".join(display_lines)
            target["activitySources"] = structured
            logger.info(
                f"  '{target['label']}': {len(display_lines)} activities extracted"
            )

    # Same claim-grounding check we run on targets, applied per-activity. An
    # activity whose `text` contains a number/year/unit not present in its own
    # `sourceText` gets a `_provenanceFlag` so reviewers can tell the difference
    # between a verbatim activity and one the LLM rewrote on the fly.
    _log_unsourced_activities(targets)

    return targets


def _log_unsourced_activities(targets: list[dict[str, Any]]) -> None:
    """Run the claim-grounding validator on each Phase 3 activity entry.

    Activities are stored on each target as a parallel array
    `activitySources: [{text, sourceText?, section?}, ...]`. Mutates each
    flagged entry in place by setting `_provenanceFlag`.
    """
    flagged = 0
    total = 0
    for target in targets:
        for entry in target.get("activitySources") or []:
            if not isinstance(entry, dict):
                continue
            total += 1
            # Reuse the same shape validate_claim_grounding expects.
            shim = {
                "text": entry.get("text", ""),
                "sources": [{"sourceText": entry.get("sourceText", "")}]
                if entry.get("sourceText")
                else [],
            }
            missing = validate_claim_grounding(shim)
            if not missing:
                continue
            flagged += 1
            entry["_provenanceFlag"] = (
                "UNSOURCED CLAIMS — the following appear in activity `text` "
                "but not in its `sourceText`: " + ", ".join(missing)
            )
            logger.warning(
                "Unsourced claim(s) in activity '%s': %s — flagged for review",
                entry.get("text", "")[:60],
                ", ".join(missing),
            )
    if flagged:
        logger.warning(
            "Validator flagged %d/%d activities with unsourced claims",
            flagged,
            total,
        )


# ---------------------------------------------------------------------------
# Main extraction pipeline
# ---------------------------------------------------------------------------


async def extract_from_text(
    text: str | DocumentText,
    doc_type: str = "NDC",
    source_document: str = "NDC",
    *,
    min_length: int = MIN_TARGET_LENGTH,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
    overlap_chars: int = OVERLAP_CHARS,
    skip_relevance_filter: bool = False,
    language: str | None = None,
) -> list[dict[str, Any]]:
    """
    Extract policy targets from document text using a multi-phase LLM approach.

    Returns a list of dicts with keys: text, label, sourceDocument, pageNumbers.
    """
    RUN_META.clear()

    # Normalize input
    if isinstance(text, str):
        doc = DocumentText(pages=[PageSpan(page=0, text=text)])
    else:
        doc = text

    full_text = doc.full_text
    if not full_text.strip():
        return []

    # Language blocks: multi-block documents (parallel translations) get
    # block-aware handling; single-block documents keep the exact legacy
    # behaviour (and therefore their LLM cache).
    blocks = split_language_blocks(doc)
    multi_block = len(blocks) > 1

    # Language detection. The working language is the explicit --language
    # override when given; for multi-block documents it is otherwise the
    # largest known-language block (the head/mid/tail majority vote is
    # meaningless when the document opens with two other languages).
    lang_code = language or "en"
    lang_name = _lang_display(language) if language else "English"
    if not language:
        known = [b for b in blocks if b.lang != "und"] if multi_block else []
        if known:
            largest = max(known, key=lambda b: b.chars)
            lang_code, lang_name = largest.lang, _lang_display(largest.lang)
        else:
            lang_code, lang_name = detect_language(full_text)

    is_english = lang_code == "en"

    # Chunking. Multi-block documents chunk per retained block so no chunk
    # straddles a language boundary; each chunk knows its block's language.
    if multi_block:
        kept_blocks = await _select_language_blocks(blocks, lang_code)
        chunks = []
        chunk_langs: dict[int, str] = {}
        for block in kept_blocks:
            block_chunks = chunk_text(
                DocumentText(pages=block.pages),
                max_chars=max_chunk_chars,
                overlap_chars=overlap_chars,
            )
            for c in block_chunks:
                chunk_langs[id(c)] = block.lang
            chunks.extend(block_chunks)
    else:
        chunks = chunk_text(doc, max_chars=max_chunk_chars, overlap_chars=overlap_chars)
        chunk_langs = {id(c): lang_code for c in chunks}
    expected = _expected_count(source_document, doc_type)
    logger.info(
        f"Extracting from {len(chunks)} chunks ({len(full_text)} chars, "
        f"language={lang_code})"
    )

    # Per-chunk accounting: one record per chunk, stamped through every phase
    # (relevance verdict, extraction outcome, candidate count) so a chunk that
    # yields nothing is diagnosable from the meta sidecar alone.
    RUN_META["language"] = {"code": lang_code, "explicit": bool(language)}
    chunk_records: list[dict[str, Any]] = [
        {
            "index": i,
            "pages": [c.pages[0], c.pages[-1]] if c.pages else [],
            "chars": len(c.text),
            "language": chunk_langs.get(id(c), lang_code),
            "relevance": "not-run",
        }
        for i, c in enumerate(chunks)
    ]
    RUN_META["chunks"] = chunk_records
    _chunk_record = {id(c): rec for c, rec in zip(chunks, chunk_records)}

    # Phase 0: Relevance filter (concurrent)
    if not skip_relevance_filter and len(chunks) > 1:
        filterable = [c for c in chunks if len(c.text.strip()) >= 80]
        for c in chunks:
            if len(c.text.strip()) < 80:
                _chunk_record[id(c)]["relevance"] = "short-skip"
        if filterable:
            filter_calls = [
                {
                    "system": RELEVANCE_FILTER_SYSTEM,
                    "user": RELEVANCE_FILTER_USER.format(
                        doc_type=doc_type,
                        # Head + middle + tail sample, NOT just the head: a
                        # chunk that opens with boilerplate may still hold
                        # the goals section further in.
                        text=_relevance_sample(chunk.text),
                    ),
                    "max_tokens": 50,
                }
                for chunk in filterable
            ]
            filter_results = await call_llm_batch_detailed(
                filter_calls,
                cache_namespace="relevance_filter",
                desc="Relevance filter",
            )
            relevant_chunks = []
            dropped_pages: list[int] = []
            for chunk, (raw, info) in zip(filterable, filter_results):
                rec = _chunk_record[id(chunk)]
                if info["status"] != "ok":
                    # A filtered/empty relevance call defaults to "relevant"
                    # in _parse_relevance — label it so that default is
                    # visible rather than silent.
                    rec["relevanceCallStatus"] = info["status"]
                if _parse_relevance(raw):
                    relevant_chunks.append(chunk)
                    rec["relevance"] = "relevant"
                else:
                    rec["relevance"] = "filtered"
                    dropped_pages.extend(chunk.pages)
                    logger.info(
                        f"  Filtered out chunk (pages {chunk.pages}): irrelevant"
                    )
            if dropped_pages:
                # Reviewer transparency: which pages the filter skipped.
                RUN_META["relevanceFilteredPages"] = sorted(set(dropped_pages))

            # Add back any very short chunks that were skipped
            short_chunks = [c for c in chunks if len(c.text.strip()) < 80]
            chunks = relevant_chunks + short_chunks
            logger.info(
                f"Relevance filter: {len(filterable)} -> {len(relevant_chunks)} "
                f"relevant chunks"
            )
    else:
        logger.info("Skipping relevance filter")

    # Filter out very short chunks
    chunks = [c for c in chunks if len(c.text.strip()) >= 80]
    if not chunks:
        return []

    # Build system prompt with optional language addendum. The addendum is
    # per chunk: in a multi-block document a kept non-working-language block
    # extracts in translate mode while working-language chunks do not.
    base_system = EXTRACT_SYSTEM.format(
        few_shot=FEW_SHOT_EXAMPLES,
        doc_type=doc_type,
        expected_count=expected,
    )

    def _system_for(c_lang: str) -> str:
        if c_lang == "en":
            return base_system
        name = lang_name if c_lang == lang_code else _lang_display(c_lang)
        return base_system + MULTILANG_ADDENDUM.format(language_name=name)

    # Phase 1: Per-chunk extraction (concurrent)
    extract_calls = [
        {
            "system": _system_for(chunk_langs.get(id(chunk), lang_code)),
            "user": EXTRACT_USER.format(doc_type=doc_type, text=chunk.text),
            "max_tokens": EXTRACT_MAX_TOKENS,
        }
        for chunk in chunks
    ]
    raw_results = await call_llm_batch_detailed(
        extract_calls,
        cache_namespace="extract",
        desc="Target extraction",
    )

    all_candidates: list[dict[str, Any]] = []
    chunk_parse_failures = 0
    for chunk, (raw, info) in zip(chunks, raw_results):
        items, parse_ok, truncated = _parse_json_array_status(
            raw, min_length=min_length, salvage_truncated=True
        )
        if not parse_ok:
            chunk_parse_failures += 1
        rec = _chunk_record.get(id(chunk))
        if rec is not None:
            rec["extraction"] = {
                "status": info["status"],
                "cached": info["cached"],
                "parseOk": parse_ok,
                "candidates": len(items),
            }
            if truncated:
                rec["extraction"]["truncated"] = True
        if truncated:
            RUN_META.setdefault("documentWarnings", []).append({
                "code": "TRUNCATED_EXTRACTION",
                "pages": chunk.pages,
                "recovered": len(items),
            })
        if info["status"] == "content_filter":
            RUN_META.setdefault("documentWarnings", []).append({
                "code": "CONTENT_FILTER",
                "phase": "extraction",
                "pages": chunk.pages,
            })
        logger.info(f"  Chunk (pages {chunk.pages}): {len(items)} candidates")

        c_lang = chunk_langs.get(id(chunk), lang_code)
        for item in items:
            item["sourceDocument"] = source_document
            item["pageNumbers"] = chunk.pages
            if c_lang != "en":
                item["language"] = c_lang
                if item.get("textOriginal"):
                    # The original came verbatim from the document; the
                    # English working text is machine-translated.
                    item["textOriginalSource"] = "source"
            all_candidates.append(item)

    logger.info(f"Phase 1 total: {len(all_candidates)} candidates")
    if chunk_parse_failures:
        RUN_META["chunkParseFailures"] = chunk_parse_failures
        logger.warning(
            "%d of %d chunk extraction responses could not be parsed — "
            "targets on those pages may be missing",
            chunk_parse_failures, len(chunks),
        )

    if len(all_candidates) == 0:
        return []

    # Deterministic dedup of chunk-overlap duplicates — always on, so small
    # documents that skip the LLM consolidation still get deduplicated.
    all_candidates = dedupe_candidates(all_candidates)

    # Phase 2: LLM consolidation in page-ordered windows (skipped only for
    # tiny candidate sets, where there is nothing meaningful to merge).
    if len(all_candidates) <= 5:
        logger.info("Skipping LLM consolidation (<=5 candidates)")
        final = all_candidates
    else:
        final = await _consolidate_windows(
            all_candidates, doc_type=doc_type, expected=expected,
            min_length=min_length,
        )
        # Cross-window duplicates (same target surfacing in two windows)
        # are merged deterministically instead of a second LLM pass.
        final = dedupe_candidates(final)

    for item in final:
        item["sourceDocument"] = source_document
        if not is_english:
            item["language"] = lang_code
            if item.get("textOriginal"):
                item["textOriginalSource"] = "source"
        # NOTE: no pageNumbers fallback. Unioning ALL candidates' pages onto
        # an item that lost its own would fabricate provenance; absent pages
        # are more honest than wrong ones.

    # Deterministic post-passes: restore document-provided labels that the
    # consolidation rephrased away, then sort into document order — reviewers
    # compare the list against the document top-to-bottom.
    _restore_doc_labels(final, all_candidates)
    final = _sort_by_document_position(final, doc)

    # Validate that synthesised text is grounded in its sources before we move
    # on to activities extraction. Logs warnings and stamps `_provenanceFlag`
    # on any target whose display text contains a number / year / unit not
    # present in any source span.
    _log_unsourced_claims(final)

    # Phase 3: Extract activities/sub-measures for each target (runs for the
    # <=5-candidate path too; the old early return skipped it).
    if final:
        final = await _extract_activities(final, doc, chunks, is_english, lang_name)

    # Quote-in-document validation: every claimed verbatim quote on targets
    # and activities must be locatable in the parsed document text. Flags,
    # never drops — the human reviewer decides.
    quotes_missing = validate_quotes_in_document(final, MatchCorpus(full_text))
    if quotes_missing:
        RUN_META["quotesNotFound"] = quotes_missing
        logger.warning(
            "Quote validator: %d claimed quote(s) could not be located in the "
            "document — flagged for review",
            quotes_missing,
        )

    return final


def _restore_original_text(
    items: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
) -> None:
    """Deterministically re-attach textOriginal lost in consolidation.

    The consolidation prompt requires textOriginal/labelOriginal to be
    preserved, but model compliance is not guaranteed (observed: one window
    dropping the field on every output). When an output item lacks
    textOriginal and its first source quote identifies exactly one input
    candidate, that candidate's original-language fields are restored.
    Synthesis items merging several candidates keep only what the model
    returned: there is no verbatim original for merged text, the sources
    array is its original-language provenance.
    """
    by_quote: dict[str, dict[str, Any]] = {}
    for cand in candidates:
        if not (cand.get("textOriginal") or cand.get("language")):
            continue
        for src in cand.get("sources") or []:
            key = normalise_for_matching(str(src.get("sourceText", "")))
            if key:
                # First writer wins; duplicate quotes across candidates are
                # ambiguous and skipped at lookup time via sentinel.
                by_quote[key] = cand if key not in by_quote else {}
    restored = 0
    for item in items:
        sources = item.get("sources") or []
        if not sources:
            continue
        key = normalise_for_matching(str(sources[0].get("sourceText", "")))
        cand = by_quote.get(key)
        if not cand:
            continue
        if (
            not item.get("textOriginal")
            and cand.get("textOriginal")
            and item.get("textCleanup") != "synthesis"
        ):
            item["textOriginal"] = cand["textOriginal"]
            if cand.get("labelOriginal"):
                item["labelOriginal"] = cand["labelOriginal"]
            restored += 1
        # In a multi-block document the per-candidate language stamp is the
        # only record of which language block an item came from; restore it
        # the same way (monolingual documents re-stamp it later anyway).
        if not item.get("language") and cand.get("language"):
            item["language"] = cand["language"]
        if (
            item.get("textOriginal")
            and item.get("language")
            and not item.get("textOriginalSource")
        ):
            item["textOriginalSource"] = "source"
    if restored:
        logger.info(
            "Restored textOriginal on %d consolidated target(s) from their "
            "source candidates",
            restored,
        )


# A label that looks document-provided: one or two words plus numbering
# ("Policy 3", "Forestry 1", "Goal 2.3"), hierarchical numbering ("4.2.6",
# "2/3"), a letter-number code ("A7.3", "M24"), or numbering followed by a
# name ("1.1 Food/agri legislation drafts"). Bare integers are NOT accepted
# as derived labels — "1" alone says nothing without its table context.
_DOC_LABEL_RE = re.compile(
    r"^(?:"
    r"[A-Za-z][\w()-]*(?:\s+[A-Za-z][\w()-]*)?\s+\d+(?:[./]\d+)*[a-z]?"
    r"|\d+(?:[./]\d+)+[a-z]?"
    r"|[A-Z]{1,3}\d+(?:\.\d+)*"
    r"|\d+(?:[./]\d+)*[.)]?\s+\S.+"
    r")$",
    re.IGNORECASE,
)
_LEADING_NUMBERING_RE = re.compile(r"^\s*(\d+(?:[./]\d+)+[a-z]?)[.)\s|-]")


def _restore_doc_labels(
    items: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
) -> None:
    """Deterministically restore document-provided labels lost downstream.

    Phase 1 is instructed to keep document-provided names ("Policy 1",
    "Goal 2.3") as labels, but consolidation can rephrase them into invented
    noun phrases. For each final item whose label does not look
    document-provided, try, in order: (a) its first source's section anchor,
    (b) the leading hierarchical numbering of its first source quote,
    (c) the label of the unique candidate sharing that source quote (same
    ambiguity sentinel as _restore_original_text). On success the label is
    replaced and labelSource="document" recorded; otherwise the LLM label
    stays — a derived label is never invented.
    """
    by_quote: dict[str, dict[str, Any]] = {}
    for cand in candidates:
        for src in cand.get("sources") or []:
            key = normalise_for_matching(str(src.get("sourceText", "")))
            if key:
                by_quote[key] = cand if key not in by_quote else {}
    restored = 0
    for item in items:
        if _DOC_LABEL_RE.match(str(item.get("label", "")).strip()):
            item["labelSource"] = "document"
            continue
        sources = item.get("sources") or []
        if not sources:
            continue
        derived: str | None = None
        section = str(sources[0].get("section", "")).strip()
        if section and _DOC_LABEL_RE.match(section):
            derived = section
        if not derived:
            m = _LEADING_NUMBERING_RE.match(str(sources[0].get("sourceText", "")))
            if m:
                derived = m.group(1)
        if not derived:
            key = normalise_for_matching(str(sources[0].get("sourceText", "")))
            cand = by_quote.get(key)
            if cand:
                cand_label = str(cand.get("label", "")).strip()
                if _DOC_LABEL_RE.match(cand_label):
                    derived = cand_label
        if derived:
            item["label"] = derived[:80]
            item["labelSource"] = "document"
            restored += 1
    if restored:
        logger.info(
            "Restored document-provided labels on %d target(s)", restored
        )


def _sort_by_document_position(
    items: list[dict[str, Any]],
    doc: DocumentText,
) -> list[dict[str, Any]]:
    """Stable-sort the final targets into document order.

    Consolidation windows may re-emit targets in any order, and multi-block
    documents concatenate per-block extractions; reviewers compare the list
    against the document top-to-bottom. Sort key = character offset of the
    first locatable source quote in the normalised document; fallback = the
    offset of the item's first page; items with neither keep their relative
    position at the end. Stable, deterministic, no fields fabricated.
    """
    norm_doc = normalise_for_matching(doc.full_text)
    # Approximate normalised offset of each page's start (fallback key only).
    page_offset: dict[int, int] = {}
    norm_len = 0
    for span in doc.pages:
        page_offset.setdefault(span.page, norm_len)
        norm_len += len(normalise_for_matching(span.text)) + 1

    def key(indexed: tuple[int, dict[str, Any]]) -> tuple[int, int]:
        idx, item = indexed
        for src in item.get("sources") or []:
            quote = normalise_for_matching(str(src.get("sourceText", "")))
            if quote:
                pos = norm_doc.find(quote)
                if pos >= 0:
                    return (pos, idx)
        pages = item.get("pageNumbers") or []
        if pages and min(pages) in page_offset:
            return (page_offset[min(pages)], idx)
        return (len(norm_doc) + 1, idx)

    return [item for _idx, item in sorted(enumerate(items), key=key)]


async def _consolidate_windows(
    candidates: list[dict[str, Any]],
    *,
    doc_type: str,
    expected: str,
    min_length: int,
) -> list[dict[str, Any]]:
    """Run the LLM consolidation pass over windows of candidates.

    One giant call must re-emit every candidate in a single response; past
    ~25 candidates that reliably truncates mid-JSON at the completion-token
    limit and the whole document silently extracts to zero (observed on
    PNSH: 79 candidates -> 0). Windows are page-ordered, so near-duplicates
    from chunk overlap sit in the same window.

    A window whose response cannot be parsed KEEPS its un-consolidated
    candidates: Phase 1 output is never discarded silently.
    """
    windows = [
        candidates[i : i + CONSOLIDATE_WINDOW]
        for i in range(0, len(candidates), CONSOLIDATE_WINDOW)
    ]
    consolidate_sys = CONSOLIDATE_SYSTEM.format(
        doc_type=doc_type, expected_count=expected
    )
    calls = []
    for window in windows:
        candidates_json = json.dumps(
            [
                {
                    "text": c["text"],
                    "label": c["label"],
                    "pageNumbers": c.get("pageNumbers", []),
                    "sources": c.get("sources", []),
                    "textCleanup": c.get("textCleanup", "verbatim"),
                    **(
                        {"textOriginal": c["textOriginal"]}
                        if c.get("textOriginal")
                        else {}
                    ),
                    **(
                        {"labelOriginal": c["labelOriginal"]}
                        if c.get("labelOriginal")
                        else {}
                    ),
                }
                for c in window
            ],
            indent=2,
            ensure_ascii=False,
        )
        calls.append({
            "system": consolidate_sys,
            "user": CONSOLIDATE_USER.format(
                count=len(window),
                doc_type=doc_type,
                candidates_json=candidates_json,
            ),
            "max_tokens": 16000,
        })

    raw_results = await call_llm_batch_detailed(
        calls, cache_namespace="extract_consolidate", desc="Consolidation"
    )

    final: list[dict[str, Any]] = []
    fallbacks = 0
    window_records: list[dict[str, Any]] = []
    for window, (raw, info) in zip(windows, raw_results):
        # salvage_truncated stays OFF here: the parse-failure fallback keeps
        # the whole un-consolidated window, which loses nothing.
        items, parse_ok, _ = _parse_json_array_status(raw, min_length=min_length)
        record: dict[str, Any] = {
            "candidatesIn": len(window),
            "itemsOut": len(items) if parse_ok else len(window),
            "fallback": not parse_ok,
        }
        if info["status"] != "ok":
            record["callStatus"] = info["status"]
        window_records.append(record)
        if not parse_ok:
            fallbacks += 1
            logger.warning(
                "Consolidation response unparseable for a %d-candidate window "
                "(likely truncated) — keeping the un-consolidated candidates",
                len(window),
            )
            final.extend(window)
        else:
            _restore_original_text(items, window)
            final.extend(items)

    RUN_META["consolidationWindows"] = window_records
    if fallbacks:
        RUN_META["consolidationFallbacks"] = fallbacks
    logger.info(
        f"Phase 2 consolidation: {len(candidates)} -> {len(final)} targets "
        f"({len(windows)} window(s), {fallbacks} fallback(s))"
    )
    return final


def _log_unsourced_claims(targets: list[dict[str, Any]]) -> None:
    """Run the claim-grounding validator and warn / flag any unsourced claims.

    Mutates the targets in place: any target with unsourced claims gets a
    `_provenanceFlag` describing what was missing, so downstream consumers
    can surface a review badge.
    """
    flagged = 0
    for t in targets:
        missing = validate_claim_grounding(t)
        if not missing:
            continue
        flagged += 1
        label = t.get("label") or t.get("text", "")[:40]
        logger.warning(
            "Unsourced claim(s) in target '%s': %s — flagged for review",
            label,
            ", ".join(missing),
        )
        t["_provenanceFlag"] = (
            "UNSOURCED CLAIMS — the following appear in `text` but not in any "
            "`sources[].sourceText`: " + ", ".join(missing)
        )
    if flagged:
        logger.warning(
            "Validator flagged %d/%d targets with unsourced claims",
            flagged,
            len(targets),
        )


async def extract_from_file(
    file_path: Path,
    doc_type: str = "NDC",
    source_document: str = "NDC",
    *,
    min_length: int = MIN_TARGET_LENGTH,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
    overlap_chars: int = OVERLAP_CHARS,
    skip_relevance_filter: bool = False,
    language: str | None = None,
) -> list[dict[str, Any]]:
    """Extract policy targets from a supported document file."""
    logger.info(f"Extracting text from {file_path.name}...")
    doc_text = extract_text(file_path)
    if not doc_text.full_text.strip():
        logger.warning(f"No text extracted from {file_path.name}")
        return []
    return await extract_from_text(
        doc_text,
        doc_type,
        source_document,
        min_length=min_length,
        max_chunk_chars=max_chunk_chars,
        overlap_chars=overlap_chars,
        skip_relevance_filter=skip_relevance_filter,
        language=language,
    )


def _emit_structured_error(code: str, detail: dict[str, Any]) -> None:
    """Print a single machine-readable error line to stdout.

    The /api/extract route parses this on a non-zero exit to surface a
    human-readable message in the upload wizard.
    """
    print(json.dumps({"error": {"code": code, **detail}}))


async def main_async(args: argparse.Namespace) -> int:
    file_path = Path(args.file)
    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        _emit_structured_error("FILE_NOT_FOUND", {"file": str(file_path)})
        return 2

    set_language(args.output_language)

    document_warnings: list[dict[str, Any]] = []
    if file_path.suffix.lower() == ".pdf":
        stats = pdf_text_layer_stats(file_path)
        no_text = stats["pages"] > 0 and (
            stats["emptyPages"] == stats["pages"] or stats["totalChars"] < 200
        )
        if no_text:
            logger.error(
                "%s appears to be a scanned or image-based PDF: %d of %d pages "
                "have no machine-readable text layer. OCR is not supported; "
                "please provide a text-based version of the document.",
                file_path.name, stats["emptyPages"], stats["pages"],
            )
            _emit_structured_error("NO_TEXT_LAYER", stats)
            return 3
        if stats["emptyRatio"] > PDF_PARTIAL_EMPTY_RATIO:
            logger.warning(
                "%s has a partial text layer: %d of %d pages have no "
                "machine-readable text. Targets on those pages cannot be "
                "extracted.",
                file_path.name, stats["emptyPages"], stats["pages"],
            )
            document_warnings.append({"code": "PARTIAL_TEXT_LAYER", **stats})

    items = await extract_from_file(
        file_path,
        doc_type=args.doc_type,
        source_document=args.source_document,
        min_length=args.min_length,
        max_chunk_chars=args.max_chunk_chars,
        overlap_chars=args.overlap_chars,
        skip_relevance_filter=args.skip_relevance_filter,
        language=args.language,
    )
    if args.country:
        for item in items:
            item["country"] = args.country

    if args.output:
        out = Path(args.output)
    else:
        out = file_path.with_suffix(".extracted.json")

    out.write_text(json.dumps(items, indent=2, ensure_ascii=False))
    logger.info(f"Saved {len(items)} targets to {out}")

    # Run metadata sidecar: document warnings (partial text layer, mixed
    # language detection) + pipeline self-reports (consolidation fallbacks,
    # quotes not found). The API route forwards these so reviewers see them
    # next to the extracted targets.
    run_meta = dict(RUN_META)
    document_warnings.extend(run_meta.pop("documentWarnings", []))
    meta = {"documentWarnings": document_warnings, **run_meta}
    if document_warnings or run_meta:
        (out.parent / (out.name + ".meta.json")).write_text(
            json.dumps(meta, indent=2)
        )

    # Persist environmental footprint of the extraction run next to the output
    # so the API route can return it and the total can be threaded into the
    # final analysis.
    footprint = get_footprint_tracker().snapshot()
    footprint_path = out.parent / (out.name + ".footprint.json")
    footprint_path.write_text(json.dumps(footprint, indent=2))
    if footprint.get("available"):
        logger.info(
            f"Extraction footprint: {footprint['energy_wh']:.4f} Wh, "
            f"{footprint['co2_geq']:.4f} gCO2eq "
            f"({footprint['tracked_call_count']} tracked, "
            f"{footprint['cached_call_count']} cached)"
        )

    # Append extraction footprint to the shared ledger (component=extract) so it
    # accumulates on the /sustainability dashboard alongside pipeline + chat.
    # Best-effort: never fail extraction on a ledger write error.
    try:
        run_id = os.getenv("CPC_RUN_ID")
        zone = electricity_zone()
        rows = footprint.get("by_model") or [footprint]
        for row in rows:
            if not int(row.get("call_count", 0) or 0):
                continue
            append_event(
                component="extract",
                provider="openai",
                model=row.get("model") or LLM_MODEL,
                region=zone,
                run_id=run_id,
                country=None,
                call_count=int(row.get("call_count", 0) or 0),
                cached_call_count=int(row.get("cached_call_count", 0) or 0),
                energy_wh=float(row.get("energy_wh", 0) or 0),
                water_ml=float(row.get("water_ml", 0) or 0),
                co2_geq=float(row.get("co2_geq", 0) or 0),
                minerals_ugsbeq=float(row.get("minerals_ugsbeq", 0) or 0),
                source=footprint.get("source", "unavailable"),
            )
    except Exception as e:
        logger.warning(f"Could not append extraction footprint ledger row: {e}")

    return 0


def main() -> None:
    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s"
    )

    parser = argparse.ArgumentParser(
        description="Extract policy targets from documents"
    )
    parser.add_argument(
        "--file", required=True, help="Path to PDF, DOCX, or text file"
    )
    parser.add_argument(
        "--doc-type", default="NDC", help="Document type label for the prompt"
    )
    parser.add_argument(
        "--source-document",
        default="NDC",
        help="PolicyDocumentType value (NDC, NBSAP, NAP, LDN, SECTORAL, OTHER)",
    )
    parser.add_argument("--output", default=None, help="Output JSON file path")
    parser.add_argument(
        "--min-length",
        type=int,
        default=MIN_TARGET_LENGTH,
        help=f"Minimum target text length (default: {MIN_TARGET_LENGTH})",
    )
    parser.add_argument(
        "--max-chunk-chars",
        type=int,
        default=MAX_CHUNK_CHARS,
        help=f"Maximum characters per chunk (default: {MAX_CHUNK_CHARS})",
    )
    parser.add_argument(
        "--overlap-chars",
        type=int,
        default=OVERLAP_CHARS,
        help=f"Overlap characters between chunks (default: {OVERLAP_CHARS})",
    )
    parser.add_argument(
        "--skip-relevance-filter",
        action="store_true",
        help="Skip the relevance filter pre-processing step",
    )
    parser.add_argument(
        "--language",
        default=None,
        help="ISO 639-1 source-document language code (e.g. 'mn' for Mongolian). "
        "Auto-detected if not provided.",
    )
    parser.add_argument(
        "--output-language",
        default="en",
        choices=["en", "es", "mn", "fr"],
        help="ISO 639-1 code for the language LLM-generated text "
        "(activities, summaries) should be returned in. Defaults to English; "
        "set to match the user's UI locale when extracting through the upload "
        "wizard.",
    )
    parser.add_argument(
        "--country",
        default=None,
        help="Country name stamped on every extracted target (curated-corpus "
        "shape), e.g. 'Panama'. Required for output that will be promoted "
        "into a targets file.",
    )
    args = parser.parse_args()

    raise SystemExit(asyncio.run(main_async(args)))


if __name__ == "__main__":
    main()
